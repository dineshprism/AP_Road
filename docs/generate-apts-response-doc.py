from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

out_path = os.path.join(os.path.dirname(__file__), "APTS-Remediation-Response.docx")

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_cell_shading(cell, fill):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("REMEDIATION RESPONSE TO APTS SECURITY AUDIT REPORT")
run.bold = True
run.font.size = Pt(14)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Andhra Pradesh Fatal Road Accident and Scientific Investigation Portal").bold = True

for line in [
    "URL: https://rspt.prismappolice.in",
    "Report: Document ID 101439, Version V1.0, dated 06-08-2026",
    "Date: 06-08-2026",
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph("Dear APTS Audit Team,")
doc.add_paragraph(
    "Thank you for the Web Application Security Assessment Report (Document ID 101439, "
    "Version V1.0, dated 06-08-2026) for the Andhra Pradesh Fatal Road Accident and "
    "Scientific Investigation Portal hosted at https://rspt.prismappolice.in."
)
doc.add_paragraph(
    "We have reviewed all observations raised in the report and have taken necessary "
    "remedial action on the production environment. The details are furnished below "
    "for your review and re-validation."
)

doc.add_paragraph("Remediation Status", style="Heading 2")

findings = [
    ("8.1", "Node.js obsolete version", "High", "Application running on unsupported Node.js version", "Application platform upgraded to a currently supported LTS release", "Remediated"),
    ("8.2", "Privilege escalation", "High", "PRISM user could access district dashboard and submit reports without authorization", "Role-based access control enforced; users restricted to authorized functions only", "Remediated"),
    ("8.3", "Unrestricted file upload", "High", "Unsafe file uploads accepted; malicious files could be uploaded", "Upload restricted to approved file types and 5 MB limit; MIME and magic-byte validation; images re-encoded and PDFs sanitized (Ghostscript) to strip active content; files served as authenticated attachments only", "Remediated"),
    ("8.4", "Host header injection", "Medium", "Invalid Host header accepted by the application", "Host header validation implemented; only trusted domain(s) permitted", "Remediated"),
    ("8.5", "Improper token timeout", "Medium", "Session remained active after period of inactivity", "30-minute inactivity timeout enforced; inactive sessions terminated", "Remediated"),
    ("8.6", "Change password not implemented", "Medium", "No change-password facility in the application", "Intentionally not implemented as per our application requirements. Self-service password change/reset is outside the scope of this portal.", "Not applicable — by design"),
    ("8.7", "Insufficient anti-automation", "Medium", "No CAPTCHA; automated login attempts possible", "CAPTCHA and rate limiting implemented on login", "Remediated"),
    ("8.8", "Security headers not implemented", "Low", "Standard HTTP security headers missing", "Required security headers configured on application responses", "Remediated"),
    ("8.9", "Concurrent logins possible", "Low", "Same account active from multiple sessions", "Single active session per user enforced", "Remediated"),
    ("8.10", "Password autocomplete enabled", "Low", "Browser could store login credentials", "Autocomplete disabled on login form", "Remediated"),
    ("8.11", "Last login time not implemented", "Low", "Previous login time not shown to user", "Last successful login time displayed after authentication", "Remediated"),
]

headers = ["#", "Finding", "Severity", "What was reported", "Action taken", "Status"]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
for i, h in enumerate(headers):
    hdr_cells[i].text = h
    for p in hdr_cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
    set_cell_shading(hdr_cells[i], "D9E2F3")

for row in findings:
    cells = table.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_paragraph()
doc.add_paragraph("Summary", style="Heading 2")

summary_data = [
    ("Total observations", "11"),
    ("Remediated", "10"),
    ("Not applicable (by requirement)", "1 (Finding 8.6)"),
    ("High severity — remediated", "3 of 3"),
    ("Medium severity — remediated", "3 of 4"),
    ("Low severity — remediated", "4 of 4"),
]
summary = doc.add_table(rows=len(summary_data), cols=2)
summary.style = "Table Grid"
for i, (k, v) in enumerate(summary_data):
    summary.rows[i].cells[0].text = k
    summary.rows[i].cells[1].text = v

doc.add_paragraph()
doc.add_paragraph("Clarification on Finding 8.6", style="Heading 2")
doc.add_paragraph(
    "Finding 8.6 (Change password functionality) has not been implemented intentionally, "
    "as it is not part of our application requirements. The portal was designed without "
    "self-service password change or reset. User credentials are managed through our "
    "established operational process."
)
doc.add_paragraph(
    "We request that this observation be recorded as not applicable for the current application."
)

doc.add_paragraph()
doc.add_paragraph("Deployment Confirmation", style="Heading 2")
doc.add_paragraph(
    "All remedial measures listed above (except Finding 8.6, which is not applicable) have "
    "been deployed on the production environment: https://rspt.prismappolice.in"
)

doc.add_paragraph()
doc.add_paragraph("Request for Re-validation", style="Heading 2")
doc.add_paragraph(
    "We request APTS to conduct re-validation of the production application against the "
    "original test cases documented in the audit report. Test user credentials for district, "
    "PRISM, and state-level roles will be provided as required."
)
doc.add_paragraph(
    "Please confirm the schedule for closure re-validation at your convenience."
)

doc.add_paragraph()
doc.add_paragraph("Thank you for your assessment and guidance.")
doc.add_paragraph()
doc.add_paragraph("Yours faithfully,")
doc.add_paragraph()
for line in [
    "[Your Name]",
    "[Designation]",
    "[Department — Home Department / PRISM Team]",
    "[Contact Number]",
    "[Email Address]",
]:
    doc.add_paragraph(line)

doc.save(out_path)
print(out_path)
